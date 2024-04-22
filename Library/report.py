import base64
import os
from bs4 import BeautifulSoup
import getpass
import xml.etree.ElementTree as ET
from tabulate import tabulate
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import smtplib
import matplotlib.pyplot as plt
from docx.shared import Pt
from urllib.parse import unquote
from docx.shared import Inches
from pdf2docx import Converter

class report(object):

    def read_word_document_content(self, file_path):
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content
    
    def send_mail(self, sender, to, subject, body, attachments, host, email_username, email_password):
        msg = MIMEMultipart()
        msg['Subject'] = subject
        # msg['Body'] = body
        msg.attach(MIMEText(body, 'plain'))
        msg['From'] = sender
        # msg['To'] = to
        msg['To'] = ', '.join(to)  # Join multiple recipients with a comma

        for attachment in attachments:
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(attachment, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment', filename=attachment)
            msg.attach(part)

        with smtplib.SMTP(host,587) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()

            server.login(email_username,email_password)
            server.send_message(msg)

    def read_specific_data_from_html_and_write_to_file(self, url, xml_path, html_file_path, output_file_path):
            current_username = getpass.getuser()
            start_time, end_time, passed, failed, total_tc = self.read_xml(xml_path)
            percentage = (int(passed)/int(total_tc))*100
            rounded_percentage = round(percentage, 2)
            # Read the HTML file
            with open(html_file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            # Parse the HTML content
            soup = BeautifulSoup(html_content, 'html.parser')

            # Example: Extract data using Css
            time = 'a:-soup-contains("Execution Time(m):")'
            execution_time = soup.select_one(time)
            result = 'span:-soup-contains("Suite Statistics:") ~ div.col-md-12 table td.fail'
            execution_result = soup.select_one(result)

            if execution_time:
                specific_data = execution_time.get_text(strip=True)
                specific_data2 = execution_result.get_text(strip=True)
                result_text = "Result:  FAILURE" if int(specific_data2) > 0 else "Result:  SUCCESS"
                # Write to a text file
                with open(output_file_path + '.txt', 'w', encoding='utf-8') as txt_file:
                    txt_file.write(f"OVERALL SUMMARY\n\n")
                    txt_file.write(f"Triggered By:  {current_username}\n")
                    txt_file.write(f"Environment Url:  {url}\n")
                    txt_file.write(f"Start Time:  {start_time}\n")
                    txt_file.write(f"End Time:  {end_time}\n")
                    txt_file.write(f"{specific_data}\n")
                    txt_file.write(f"{result_text}\n\n\n")
                    txt_file.write(f"TEST RESULTS\n\n")
                    txt_file.write(f"Total: {total_tc}\n")
                    txt_file.write(f"Passed: {passed}\n")
                    txt_file.write(f"Failed: {failed}\n")
                    txt_file.write(f"Pass Percentage: {rounded_percentage}%\n")

                # Write to a Word document (docx)
                # doc = Document()
                # doc.add_paragraph(f"OVERALL SUMMARY\n")
                # doc.add_paragraph(f"Triggered By:  {current_username}")
                # doc.add_paragraph(f"Environment Url:  {url}")
                # doc.add_paragraph(f"Start Time:  {start_time}")
                # doc.add_paragraph(f"End Time:  {end_time}")
                # doc.add_paragraph(f"{specific_data}")
                # doc.add_paragraph(f"{result_text}\n\n")
                # doc.add_paragraph(f"TEST RESULTS\n")
                # doc.add_paragraph(f"Total: {total_tc}")
                # doc.add_paragraph(f"Passed: {passed}")
                # doc.add_paragraph(f"Failed: {failed}")
                # doc.add_paragraph(f"Pass Percentage: {rounded_percentage}%")
                # doc.save(output_file_path + '.docx')

            else:
                print("Specific element not found.")
        
    def read_specific_data(self, url, html_file_path):
        current_username = getpass.getuser()
        print(f"Triggered By: {current_username}")
        print(f"Environment: " + url)
        # Read the HTML file
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        # Example: Extract data using XPath
        time = 'a:-soup-contains("Execution Time(m):")'
        execution_time = soup.select_one(time)
        result = 'span:-soup-contains("Suite Statistics:") ~ div.col-md-12 table td.fail'
        execution_result = soup.select_one(result)
        table_selector = 'table#tm'
        execution_table = soup.select_one(table_selector)

        specific_data = execution_time.get_text(strip=True)
        print(specific_data)
        specific_data2 = execution_result.get_text(strip=True)
        if int(specific_data2) > 0:
            print("Result: FAILURE")
        else:
            print("Result: SUCCESS")
            
        # Extract data from the table in a tabular format
        table_data = []
        for row in execution_table.select('tr'):
            row_data = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
            table_data.append(row_data)
        # Print the table using tabulate
        print(tabulate(table_data, headers="firstrow", tablefmt="grid"))

    def read_html_and_generate_text(self, url, html_file_path, output_text_path):
        current_username = getpass.getuser()

        # Read the HTML file
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')

        # Example: Extract data using XPath
        time = 'a:-soup-contains("Execution Time(m):")'
        execution_time = soup.select_one(time)
        result = 'span:-soup-contains("Suite Statistics:") ~ div.col-md-12 table td.fail'
        execution_result = soup.select_one(result)
        testcase_results = 'table#tm'
        testcase_table = soup.select_one(testcase_results)
        suite = 'table#sm'
        suite_table = soup.select_one(suite)
        svg_tags = soup.find_all('svg', {'class': 'apexcharts-svg'})
        for svg in svg_tags:
            print(svg)

        if execution_time and testcase_table:
            specific_data = execution_time.get_text(strip=True)
            specific_data2 = execution_result.get_text(strip=True)

            # Write print statements to a text file
            with open(output_text_path, 'w', encoding='utf-8') as text_file:
                text_file.write(f"*Overall Summary*\n\n")
                # text_file.write(f"\n")
                text_file.write(f"Name: Regression Run\n")
                text_file.write(f"Triggered By: {current_username}\n")
                text_file.write(f"Environment: {url}\n")
                text_file.write(f"{specific_data}\n")
                text_file.write(f"Result: {'FAILURE' if int(specific_data2) > 0 else 'SUCCESS'}\n")

                # Write the table to the text file
                text_file.write("\n\n*Suite Statistics*\n")
                table_data = []
                for row in suite_table.select('tr'):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(row_data)
                table_text = tabulate(table_data, headers="firstrow", tablefmt="grid")
                text_file.write(table_text)

                text_file.write("\n\n\n*Test Statistics*\n")
                table_data = []
                for row in testcase_table.select('tr'):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(row_data)
                table_text = tabulate(table_data, headers="firstrow", tablefmt="grid")
                text_file.write(table_text)
        else:
            print("Specific element not found.")

    def pdf_to_word(self, input_pdf, output_word):
        cv = Converter(input_pdf)
        cv.convert(output_word, start=0, end=None)
        cv.close()
    
    def read_xml(self, input_xml_path):
            # Parse the XML data from the file
            tree = ET.parse(input_xml_path)
            root = tree.getroot()
            status_elements = root.findall(".//test/status")
            # Initialize a counter for 'SKIP'
            skip_count = 0
            # Iterate through the status elements
            for status_element in status_elements:
                status_value = status_element.get('status')
                if status_value == 'SKIP':
                    skip_count += 1
            # Print the count of 'SKIP'
            print("Number of 'SKIP' statuses:", skip_count)        
            suite_with_stat = root.find(".//suite/status")
            if suite_with_stat is not None:
                starttime_value = suite_with_stat.get('starttime')
                endtime_value = suite_with_stat.get('endtime')
            statistics = root.find(".//statistics/total/stat")
            passed = statistics.get('pass')
            failed = statistics.get('fail')
            total_tc = int(passed)+int(failed)+int(skip_count)
            return starttime_value, endtime_value, passed, failed, skip_count, total_tc

    def seconds_to_minutes_and_seconds(self, seconds):
        seconds = float(seconds)  # Convert to float (or int) if it's a string
        minutes = int(seconds // 60)
        remaining_seconds = int(seconds % 60)
        return f"{minutes} min {remaining_seconds} sec"

    def read_html_xml_and_write_html(self, url, xml_path, html_file_path, output_path, kw_details='false'):
        current_username = getpass.getuser()
        # Read the Output.XML file
        start_time, end_time, passed, failed, skipped, total_tc = self.read_xml(xml_path)
        actual_total = int(passed)+int(failed)
        percentage = (int(passed)/int(actual_total))*100
        rounded_percentage = round(percentage, 2)
        splitted_xml_name = xml_path.split("-")
        xml_time = splitted_xml_name[2].split(".")[0]
        xml_hr = xml_time[:2]
        build_no = splitted_xml_name[1]+"."+xml_hr
        output_html_path = output_path + build_no+".html"
        # Read the robotmetrics.HTML file
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        # Example: Extract data using CSS
        time = 'a:-soup-contains("Execution Time(m):")'
        execution_time = soup.select_one(time)
        result = 'span:-soup-contains("Suite Statistics:") ~ div.col-md-12 table td.fail'
        execution_result = soup.select_one(result)
        testcase_results = 'table#tm tbody'
        testcase_table = soup.select_one(testcase_results)
        suite = 'table#sm tbody'
        suite_table = soup.select_one(suite)
        keywords = 'table#kmt tbody'
        kw_table = soup.select_one(keywords)
        kw_table_tr = soup.select_one('table#kmt tbody tr')

        # Pie chart, where the slices will be ordered and plotted counter-clockwise:
        labels = 'Fail', 'Pass'
        sizes = [failed, passed]
        colors = ['red', 'green']
        explode = (0, 0.02)
        # Plotting the pie chart
        plt.figure(figsize=(3, 3), facecolor='none')
        plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90, wedgeprops=dict(width=0.4, edgecolor='w'))
        plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        # Save the plot as an image file
        plt.savefig('chart.png', transparent=True)
        with open('chart.png', 'rb') as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        with open('Logo.png', 'rb') as image_file1:
            base64_image1 = base64.b64encode(image_file1.read()).decode('utf-8')

        if execution_time and testcase_table:
            specific_data = execution_time.get_text(strip=True)
            execution, time = specific_data.split(":")
            specific_data2 = execution_result.get_text(strip=True)
            result_text = 'FAILURE' if int(specific_data2) > 0 else 'SUCCESS'
            result_color = 'failure' if int(specific_data2) > 0 else 'success'
            with open(output_html_path, 'w', encoding='utf-8') as html_file:
                # Write print statements to a html file
                html_content = f"<html>\n<head>\n<style>\n"
                # html_content += "body { background-color: #f0f0f0; }\n"
                html_content += "strong.keys { font-weight: bold; color: black; }\n"
                html_content += "strong.header { font-weight: bold; font-size: 30px; color: green; }\n"
                html_content += "strong.glow { font-weight: bold; color: black; animation: glow 1s ease-in-out infinite alternate; }\n"
                html_content += "@keyframes glow { to { text-shadow: 0 0 20px #00fffc, 0 0 60px #00fffc; } }\n"
                html_content += "strong { font-weight: bold; color: blue; }\n"
                html_content += "strong.result.success { color: green; }"
                html_content += "strong.result.failure { color: red; }"
                html_content += "table { border-collapse: collapse; border: 1px solid black; width: 100%; }\n"
                html_content += "th { border: 1px solid black; text-align: center; padding: 8px; }\n"
                html_content += "td { border: 1px solid black; text-align: left; padding: 8px; }\n"
                html_content += "</style>\n</head>\n<body>"
                html_content += "<div style='display: flex;'>"
                html_content += "    <div style='flex: 1;'>"
                html_content += "       <div style='text-align: center;'><strong class='header'>VLMS Automation Report</strong></div>"
                html_content += "    </div>"
                html_content += "    <div style='flex: 1; text-align: right;'>"
                html_content += f"       <h><img src='data:image/png;base64,{base64_image1}' alt='Vg_logo'></h>"
                html_content += "    </div>"
                html_content += "</div>"                
                html_content += "<br>\n\n\n"
                html_content += "<div style='display: flex;'>"
                html_content += "    <div style='flex: 0.5;'>"
                html_content += "        <p>\n<strong class='keys glow'>Plan Report:</strong></p>"
                html_content += "        <p><strong>Name:</strong> Regression Test Suites</p>\n"
                html_content += f"        <p><strong>Triggered By:</strong> {current_username}</p>\n"
                html_content += f"        <p><strong>Environment Url:</strong> {url}</p>\n"
                html_content += f"        <p><strong>Start Time:</strong> {start_time}</p>\n"
                html_content += f"        <p><strong>End Time:</strong> {end_time}</p>\n"
                html_content += f"        <p><strong>{execution}:</strong> {time}</p>\n"
                html_content += f"        <p><strong>Result:</strong> <strong class='result {result_color}'>{result_text}</strong></p>\n"
                html_content += f"        <p><strong>Build Number:</strong> {build_no}</p>\n "
                html_content += "    </div>"
                html_content += "    <div style='flex: 0.3;'>"
                html_content += "<br>\n"
                # html_content += "        <h><strong class='keys'>Test Case Results</strong></h><p><img src='chart.png' alt='Test Case Pie Chart'></p>"
                html_content += f"       <h><strong class='keys glow'>Test Case Results</strong></h><p><img src='data:image/png;base64,{base64_image}' alt='Test Case Pie Chart'></p>"
                html_content += "    </div>"
                html_content += "</div>"
                html_content += "<p><strong class='keys glow'>Overall Summary:</strong></p>\n\n"
                html_content += "<table style='width: 40%;'>"
                html_content += "<tr><th style='width: 20%;'>Statistic</th><th style='width: 10%;'>Value</th></tr>"
                html_content += f"<tr><td><strong>Testcases Total:</strong></td><td>{total_tc}</td></tr>"
                html_content += f"<tr><td><strong>Testcases Passed:</strong></td><td>{passed}</td></tr>"
                html_content += f"<tr><td><strong>Testcases Failed:</strong></td><td>{failed}</td></tr>"
                html_content += f"<tr><td><strong>Testcases Retried:</strong></td><td>{skipped}</td></tr>"
                html_content += f"<tr><td><strong>Pass Percentage:</strong></td><td><strong class='result success'>{rounded_percentage}%</strong></td></tr>"
                html_content += "</table>\n"
                html_content += "<br>\n\n\n"
                html_content += "<br>\n\n\n"
                # Write the table for Test and Suite statistics to the HTML file
                html_content += "<p><strong class='keys glow'>Suite Statistics:</strong></p>\n"
                html_content += "<table style='width: 80%;'>"
                html_content += "\n<tr><th>Sl. No</th><th style='width: 30%;'>Suite Name</th><th>Status</th><th>Total</th><th>Pass</th><th>Fail</th><th>Retry</th><th style='width: 25%;'>Time</th></tr>\n"
                for index, row in enumerate(suite_table.select('tr')):
                    time = row.select_one('td:nth-of-type(7)')
                    time_in_seconds = time.get_text(strip=True)
                    row_data = []
                    html_content += "<tr>"
                    html_content += f"<td>{index + 1}</td>"
                    for col_index, cell in enumerate(row.find_all(['th', 'td'])):
                        cell_text = cell.get_text(strip=True)
                        cell_style = cell.get('style', '')
                        if col_index == 6:
                            # Convert and append time in the fourth column
                            time_in_seconds = cell_text
                            cell_text = self.seconds_to_minutes_and_seconds(time_in_seconds)
                        html_content += f"<td style='{cell_style}'>{cell_text}</td>"
                        row_data.append((cell_text, cell_style))
                    html_content += "</tr>\n"
                html_content += "</table>\n"
                html_content += "<br>\n\n"
                html_content += "<br>\n\n"
                html_content += "<p><strong class='keys glow'>Test Statistics:</strong></p>\n"
                html_content += "<table>\n<tr><th>Sl. No</th><th>Suite Name</th><th style='width: 30%;'>Test Name</th><th>Status</th><th>Time</th><th style='width: 30%;'>Error Message</th></tr>\n"
                for index, row in enumerate(testcase_table.select('tr')):
                    time = row.select_one('td:nth-of-type(4)')
                    time_in_seconds = time.get_text(strip=True)
                    row_data = []
                    html_content += "<tr>"
                    html_content += f"<td>{index + 1}</td>"
                    for col_index, cell in enumerate(row.find_all(['th', 'td'])):
                        if 'hide' in cell.get('class', []):
                            continue
                        cell_text = cell.get_text(strip=True)
                        cell_style = cell.get('style', '')
                        if col_index == 3:
                            # Convert and append time in the fourth column
                            time_in_seconds = cell_text
                            cell_text = self.seconds_to_minutes_and_seconds(time_in_seconds)
                        html_content += f"<td style='{cell_style}'>{cell_text}</td>"
                        row_data.append((cell_text, cell_style))
                    html_content += "</tr>\n"
                html_content += "</table>\n"
                html_content += "<br>\n\n"
                html_content += "<br>\n\n"
                if kw_table_tr and kw_details != 'false':
                    html_content += "<p><strong class='keys glow'>Keyword Statistics:</strong></p>\n"
                    html_content += "<table style='width: 80%;'>"
                    html_content += "\n<tr><th>Sl. No</th><th style='width: 40%;'>Keyword Name</th><th>Times</th><th style='width: 15%;'>Min Duration(s)</th><th style='width: 15%;'>Max Duration(s)</th><th style='width: 15%;'>Average Duration(s)</th></tr>\n"
                    for index, row in enumerate(kw_table.select('tr')):
                        row_data = []
                        html_content += "<tr>"
                        html_content += f"<td>{index + 1}</td>"
                        for cell in row.find_all(['th', 'td']):
                            cell_text = cell.get_text(strip=True)
                            cell_style = cell.get('style', '')
                            html_content += f"<td style='{cell_style}'>{cell_text}</td>"
                            row_data.append((cell_text, cell_style))
                        html_content += "</tr>\n"
                    html_content += "</table>\n"

                html_content += "</body>\n</html>"

                html_file.write(html_content)
        else:
            print("Specific element not found.")
        os.remove('chart.png')

    def read_xmll(self, xml_path, output_xml_path):
        # Reading the data inside the xml
        # file to a variable under the name 
        # data
        with open(xml_path, 'r') as f:
            data = f.read()
        
        # Passing the stored data inside
        # the beautifulsoup parser, storing
        # the returned object 
        Bs_data = BeautifulSoup(data, "xml")
        
        # Finding all instances of tag 
        # `unique`
        b_unique = Bs_data.find_all('.//suite/stat')
        
        print(b_unique)
        with open(output_xml_path, 'w') as output_file:
            output_file.write(b_unique.prettify())

    def read_xml_data(self, xml_path, output_xml_path):
            # Reading the data inside the xml
            # file to a variable under the name 
            # data
            with open(xml_path, 'r') as f:
                data = f.read()
            
            # Passing the stored data inside
            # the BeautifulSoup parser, storing
            # the returned object 
            soup = BeautifulSoup(data, "xml")
            
            # Finding all instances of the tag `stat` inside `suite`
            suite_stats = soup.find_all('suite/stat')

            # Print the results
            for stat_element in suite_stats:
                print(stat_element.text)

            # # Write the output to a new XML file
            with open(output_xml_path, 'w') as output_file:
                output_file.write(suite_stats.prettify())
    
    def filter_and_write_xmll(self, input_xml_path, output_xml_path):
        # Parse the XML data from the file
        tree = ET.parse(input_xml_path)
        root = tree.getroot()

        # Find all <suite> elements with <stat> children
        suites_with_stats = root.findall(".//suite[stat]")

        # Create a new XML root for the filtered data
        filtered_root = ET.Element("robot")

        # Add the filtered <suite> elements to the new root
        filtered_root.extend(suites_with_stats)

        # Create a new ElementTree for the filtered data
        filtered_tree = ET.ElementTree(filtered_root)

        # Write the filtered data to a new XML file
        filtered_tree.write(output_xml_path)

    def xml_to_html(self, xml_file, html_file):
        # Parse the XML file
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # Create an HTML document using BeautifulSoup
        soup = BeautifulSoup(features='html.parser')
        html_tag = soup.new_tag('html')
        soup.append(html_tag)

        # Create head and body tags
        head_tag = soup.new_tag('head')
        body_tag = soup.new_tag('body')
        html_tag.append(head_tag)
        html_tag.append(body_tag)

        # Add a title to the head
        title_tag = soup.new_tag('title')
        title_tag.string = 'Robot Framework Report'
        head_tag.append(title_tag)

        # Extract statistics data
        statistics = root.find('.//statistics')
        if statistics:
            # Create an HTML table
            table_tag = soup.new_tag('table')
            body_tag.append(table_tag)

            # Add table headers
            headers = ['ID', 'Name', 'Pass', 'Fail', 'Skip']
            tr_tag = soup.new_tag('tr')
            table_tag.append(tr_tag)
            for header in headers:
                th_tag = soup.new_tag('th')
                th_tag.string = header
                tr_tag.append(th_tag)

            # Add 'total' row
            total_stat = statistics.find('.//total/stat')
            if total_stat:
                tr_tag = soup.new_tag('tr')
                table_tag.append(tr_tag)
                for key in ['id', 'name', 'pass', 'fail', 'skip']:
                    td_tag = soup.new_tag('td')
                    td_tag.string = total_stat.get(key, '')
                    tr_tag.append(td_tag)

            # Add 'suite' row
            suite_stat = statistics.find('.//suite/stat')
            if suite_stat:
                tr_tag = soup.new_tag('tr')
                table_tag.append(tr_tag)
                for key in ['id', 'name', 'pass', 'fail', 'skip']:
                    td_tag = soup.new_tag('td')
                    td_tag.string = suite_stat.get(key, '')
                    tr_tag.append(td_tag)

        # Write the HTML to a file
        with open(html_file, 'w', encoding='utf-8') as html_output:
            html_output.write(str(soup))

    def handle_paragraph(self, paragraph, doc):
        # Handle paragraphs
        if paragraph.name == 'p':
            new_paragraph = doc.add_paragraph(paragraph.get_text(strip=True))

            # Check if the paragraph has runs before accessing the first run
            if new_paragraph.runs:
                # You can customize the style, font size, etc.
                run = new_paragraph.runs[0]
                run.font.size = Pt(12)

    def handle_image(self, image, doc):
        # Handle images
        if image.name == 'img':
            # Extract the image source from the HTML
            image_source = unquote(image['src'])
            # You may need to handle image downloading or embedding here
            # Example: doc.add_picture('path/to/image.jpg', width=Inches(1.0))
            doc.add_picture(image_source, width=Inches(2.0))
    
    def handle_table(self, table, doc):
        # Handle tables
        if table.name == 'table':
            new_table = doc.add_table(rows=1, cols=1)
            for row in table.find_all('tr'):
                new_row = new_table.add_row()
                for cell in row.find_all(['th', 'td']):
                    new_cell = new_row.cells[0]  # Add a new cell to the right
                    new_cell.text = cell.get_text(strip=True)

    def html_file_to_docx(self, html_path, output_path='output.docx'):
        # Read the HTML file
        with open(html_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        # Create a new Word document
        doc = Document()
        # Iterate through HTML elements and convert them to corresponding Word elements
        for element in soup.descendants:
            if element.name == 'p':
                self.handle_paragraph(element, doc)
            elif element.name == 'img':
                self.handle_image(element, doc)
            elif element.name == 'table':
                self.handle_table(element, doc)
            # Add more conditions based on your HTML structure
        # Save the Word document
        doc.save(output_path)
